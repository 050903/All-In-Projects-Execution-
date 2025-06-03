import os
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import platform
from docx import Document
# from docx.shared import Inches # Không cần thiết cho logic hiện tại
# from docx.enum.shape import WD_INLINE_SHAPE # Vẫn có thể hữu ích nếu muốn xử lý inline shapes đặc biệt
from docx.image.exceptions import UnrecognizedImageError # Cần thiết cho try-except
import subprocess # Cho việc mở VS Code

def extract_images(document, output_image_dir):
    """
    Trích xuất tất cả hình ảnh từ tài liệu Word và lưu chúng vào output_image_dir.
    Trả về một danh sách các tên tệp hình ảnh (tương đối với output_image_dir).
    """
    images = []
    image_count = 0 # Khởi tạo image_count ở đây
    
    os.makedirs(output_image_dir, exist_ok=True)

    for rel_id, rel in document.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_part = rel.target_part
                image_data = image_part.blob
                image_extension = "png" # Mặc định an toàn

                # Cố gắng lấy phần mở rộng tệp
                if hasattr(image_part, 'default_image_ext'):
                    try:
                        image_extension = image_part.default_image_ext
                    except UnrecognizedImageError:
                        print(f"Cảnh báo (rel_id: {rel_id}): default_image_ext không nhận dạng được loại ảnh, thử content_type.")
                        # Sẽ rơi vào khối content_type bên dưới nếu image_extension vẫn là "png"
                    except Exception as e_die: 
                        print(f"Cảnh báo (rel_id: {rel_id}): Lỗi với default_image_ext: {e_die}, thử content_type.")
                        # Sẽ rơi vào khối content_type bên dưới nếu image_extension vẫn là "png"
                
                # Nếu default_image_ext không có hoặc không thành công (hoặc trả về png và muốn kiểm tra content_type)
                # sử dụng content_type làm fallback hoặc để có thông tin chính xác hơn.
                if (image_extension == "png" or not hasattr(image_part, 'default_image_ext')) and hasattr(image_part, 'content_type'):
                    content_type = image_part.content_type.lower()
                    if 'jpeg' in content_type or 'jpg' in content_type:
                        image_extension = 'jpg'
                    elif 'png' in content_type: # Giữ nguyên nếu là png
                        image_extension = 'png'
                    elif 'gif' in content_type:
                        image_extension = 'gif'
                    elif 'bmp' in content_type:
                        image_extension = 'bmp'
                    elif 'tiff' in content_type:
                        image_extension = 'tiff'
                    elif 'svg+xml' in content_type:
                        image_extension = 'svg'
                    elif 'x-wmf' in content_type: 
                        image_extension = 'wmf'
                    elif 'x-emf' in content_type: 
                        image_extension = 'emf'
                    elif image_extension == "png": # Nếu vẫn là png sau khi kiểm tra content_type không khớp
                        print(f"Cảnh báo (rel_id: {rel_id}): Không xác định được phần mở rộng từ content_type: {image_part.content_type}, sử dụng '.{image_extension}'.")

                img_filename = f'image_{image_count}.{image_extension}'
                img_path_abs = os.path.join(output_image_dir, img_filename)
                
                with open(img_path_abs, 'wb') as f:
                    f.write(image_data)
                
                images.append(img_filename) 
                image_count += 1 # Tăng image_count sau khi xử lý thành công
            
            except AttributeError as ae:
                if "'ImagePart' object has no attribute 'default_image_ext'" in str(ae):
                    # Thông báo này chỉ để biết, vì logic fallback đã được áp dụng
                    print(f"Thông tin (rel_id: {rel_id}): Phiên bản python-docx có thể cũ, không có 'default_image_ext'. Đã sử dụng content_type làm fallback.")
                    # Không cần làm gì thêm ở đây vì logic fallback đã chạy
                else:
                    print(f"Lỗi AttributeError khi trích xuất hình ảnh {image_count} (rel_id: {rel_id}): {str(ae)}")
                # Không tăng image_count nếu có lỗi ở đây và không lưu được file
                continue # Bỏ qua hình ảnh này nếu có lỗi không mong muốn
            except Exception as e: 
                print(f"Lỗi chung khi trích xuất hình ảnh {image_count} (rel_id: {rel_id}): {str(e)}")
                continue # Bỏ qua hình ảnh này
    return images

def convert_word_to_latex_simple(docx_path, output_tex_path):
    try:
        document = Document(docx_path)
    except Exception as e:
        messagebox.showerror("Lỗi", f"Lỗi khi đọc tệp Word: {e}")
        return False

    latex_content = []
    
    output_dir = os.path.dirname(output_tex_path)
    if not output_dir: 
        output_dir = "." 
    
    extracted_images = extract_images(document, output_dir)

    latex_content.extend([
        "\\documentclass[12pt,a4paper]{article}",
        "\\usepackage[utf8]{inputenc}", 
        "\\usepackage[T5,T1]{fontenc}",   
        "\\usepackage[vietnamese]{babel}", 
        "\\usepackage{graphicx}",
        "\\usepackage[left=2cm,right=2cm,top=2cm,bottom=2cm]{geometry}",
        "\\usepackage{amsmath}", 
        "\\usepackage{amsfonts}",
        "\\usepackage{amssymb}",
        "\\usepackage{hyperref}", 
        "\\usepackage{svg}", # Thêm gói svg nếu bạn có hình ảnh svg
        "\\begin{document}",
        ""
    ])
    
    for paragraph in document.paragraphs:
        latex_line = ""
        is_heading = False
        heading_level_command = ""

        if paragraph.style.name.startswith('Heading'):
            try:
                level = int(paragraph.style.name.split(' ')[-1])
                if level == 1: heading_level_command = "\\section*{"
                elif level == 2: heading_level_command = "\\subsection*{"
                elif level == 3: heading_level_command = "\\subsubsection*{"
                else: heading_level_command = "\\paragraph*{"
                is_heading = True
                latex_line += heading_level_command
            except ValueError: 
                pass # Không phải định dạng "Heading X", xử lý như văn bản thường

        for run in paragraph.runs:
            text = run.text
            # Escape các ký tự đặc biệt của LaTeX
            text = text.replace('\\', '\\textbackslash{}') # Phải escape \ trước
            text = text.replace('&', '\\&')
            text = text.replace('%', '\\%')
            text = text.replace('$', '\\$')
            text = text.replace('#', '\\#')
            text = text.replace('_', '\\_')
            text = text.replace('{', '\\{')
            text = text.replace('}', '\\}')
            text = text.replace('~', '\\textasciitilde{}')
            text = text.replace('^', '\\textasciicircum{}')

            if run.bold:
                text = f"\\textbf{{{text}}}"
            if run.italic:
                text = f"\\textit{{{text}}}"
            latex_line += text
        
        if is_heading:
            # Đảm bảo heading có nội dung hoặc đóng đúng cách
            current_text_in_heading = latex_line[len(heading_level_command):].strip()
            if not current_text_in_heading: # Nếu không có text sau lệnh heading
                 latex_line += "Tiêu đề trống}" 
            else:
                 latex_line += "}" # Đóng ngoặc cho section/subsection
            latex_content.append(latex_line + "\n")
        elif latex_line.strip(): 
            latex_content.append(latex_line + "\n")
        elif not latex_content or (latex_content and latex_content[-1].strip()): # Thêm dòng trống nếu dòng trước có nội dung
             latex_content.append("") # Thêm một dòng trống để giữ khoảng cách giữa các đoạn

    if extracted_images:
        latex_content.append("\n\\clearpage % Đảm bảo hình ảnh ở trang mới nếu cần\n")
        latex_content.append("\\section*{Hình ảnh}\n")
        latex_content.append("% Các hình ảnh được trích xuất từ tài liệu Word:\n")
        for img_filename in extracted_images:
            img_path_for_latex = img_filename.replace("\\", "/")
            img_basename, img_ext = os.path.splitext(img_path_for_latex)
            
            figure_content = [
                "\\begin{figure}[htbp]", 
                "\\centering"
            ]
            if img_ext.lower() == '.svg':
                # Gói svg yêu cầu tên file không có phần mở rộng .svg
                figure_content.append(f"\\includesvg[width=0.8\\textwidth]{{{img_basename}}}") 
            else:
                figure_content.append(f"\\includegraphics[width=0.8\\textwidth]{{{img_path_for_latex}}}")
            
            figure_content.extend([
                f"\\caption{{Hình ảnh từ tài liệu: {os.path.basename(img_filename)}}}", 
                f"\\label{{fig:{os.path.splitext(os.path.basename(img_filename))[0]}}}",
                "\\end{figure}\n"
            ])
            latex_content.extend(figure_content)

    latex_content.append("\\end{document}")

    try:
        with open(output_tex_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(latex_content))
        # Thông báo về hình ảnh sẽ được xử lý trong GUI
        return True
    except Exception as e:
        messagebox.showerror("Lỗi", f"Lỗi khi ghi tệp LaTeX: {e}")
        return False

class WordToLaTeXConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Chuyển đổi Word sang LaTeX")
        self.root.geometry("800x600")

        tk.Label(root, text="Tệp Word đầu vào:").pack(pady=5)
        self.input_frame = tk.Frame(root)
        self.input_frame.pack(fill=tk.X, padx=5)

        self.input_path_var = tk.StringVar()
        tk.Entry(self.input_frame, textvariable=self.input_path_var, width=70).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        tk.Button(self.input_frame, text="Chọn tệp...", command=self.browse_input).pack(side=tk.LEFT)

        tk.Label(root, text="Tệp LaTeX đầu ra:").pack(pady=5)
        self.output_frame = tk.Frame(root)
        self.output_frame.pack(fill=tk.X, padx=5)

        self.output_path_var = tk.StringVar()
        tk.Entry(self.output_frame, textvariable=self.output_path_var, width=70).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        tk.Button(self.output_frame, text="Chọn tệp...", command=self.browse_output).pack(side=tk.LEFT)

        tk.Button(root, text="Chuyển đổi", command=self.convert).pack(pady=10)

        tk.Label(root, text="Nhật ký:").pack(pady=5)
        self.log_area = scrolledtext.ScrolledText(root, width=90, height=20, wrap=tk.WORD)
        self.log_area.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)

    def log(self, message):
        self.log_area.insert(tk.END, message + "\n")
        self.log_area.see(tk.END)
        self.root.update_idletasks()

    def browse_input(self):
        filename = filedialog.askopenfilename(
            title="Chọn tệp Word",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.input_path_var.set(filename)
            output_filename = os.path.splitext(filename)[0] + ".tex"
            self.output_path_var.set(output_filename)

    def browse_output(self):
        initial_dir = os.path.dirname(self.input_path_var.get())
        initial_file = os.path.basename(self.output_path_var.get())
        
        filename = filedialog.asksaveasfilename(
            title="Lưu tệp LaTeX",
            initialdir=initial_dir if initial_dir else os.getcwd(),
            initialfile=initial_file if initial_file else "output.tex",
            filetypes=[("LaTeX files", "*.tex"), ("All files", "*.*")],
            defaultextension=".tex"
        )
        if filename:
            self.output_path_var.set(filename)

    def convert(self):
        input_path = self.input_path_var.get()
        output_path = self.output_path_var.get()

        if not input_path or not output_path:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cả tệp đầu vào và đầu ra!")
            return

        self.log(f"Bắt đầu chuyển đổi từ '{input_path}' sang '{output_path}'...")
        
        # Chuyển hướng print để log vào GUI
        original_print = print
        def gui_print(*args, **kwargs):
            self.log(" ".join(map(str,args)))
            # original_print(*args, **kwargs) # Bỏ comment nếu vẫn muốn print ra console
        
        # Tạm thời thay thế print()
        # import builtins # Không nên thay đổi builtins trực tiếp nếu không quá cần thiết
        # old_print = builtins.print
        # builtins.print = gui_print
        # Thay vào đó, truyền hàm log vào nếu cần, hoặc để các hàm con tự print
        # và người dùng xem console. Trong trường hợp này, extract_images đã print lỗi.

        success = convert_word_to_latex_simple(input_path, output_path)
        
        # Khôi phục hàm print ban đầu
        # builtins.print = old_print

        if success:
            self.log("Chuyển đổi thành công!")
            self.log(f"Tệp LaTeX đã được lưu tại: {output_path}")
            
            output_dir_abs = os.path.abspath(os.path.dirname(output_path))
            
            user_choice = messagebox.askyesnocancel(
                "Thành công & Mở tệp",
                f"Chuyển đổi hoàn tất.\n"
                f"Tệp LaTeX: {output_path}\n"
                f"Hình ảnh (nếu có) đã được lưu vào:\n{output_dir_abs}\n\n"
                f"Bạn có muốn mở tệp LaTeX trong VS Code không?\n"
                f"(Chọn 'Yes' để mở trong VS Code, 'No' để chỉ mở thư mục chứa tệp, 'Cancel' để không làm gì.)\n\n"
                f"Lưu ý: Cần VS Code, extension LaTeX Workshop đã cài đặt và lệnh 'code' trong PATH.",
                icon='question'
            )

            if user_choice is True: 
                try:
                    cmd_args = ['code', output_path]
                    use_shell = platform.system() == "Windows"
                    subprocess.run(cmd_args, shell=use_shell, check=False)
                    self.log(f"Đã yêu cầu mở {output_path} trong VS Code.")
                    self.log("Nếu VS Code không mở, hãy đảm bảo 'code' command đã được thêm vào PATH.")
                except FileNotFoundError:
                    self.log("Lệnh 'code' (VS Code) không tìm thấy. Hãy đảm bảo VS Code đã được cài đặt và 'code' command đã được thêm vào PATH hệ thống.")
                    messagebox.showerror("Lỗi", "Lệnh 'code' (VS Code) không tìm thấy. Vui lòng mở tệp thủ công.")
                except Exception as e:
                    self.log(f"Lỗi khi cố gắng mở tệp trong VS Code: {e}")
                    messagebox.showerror("Lỗi", f"Lỗi khi mở tệp trong VS Code: {e}")
            
            elif user_choice is False: 
                try:
                    if platform.system() == "Windows":
                        os.startfile(output_dir_abs) 
                    elif platform.system() == "Darwin": 
                        subprocess.run(['open', output_dir_abs], check=True)
                    else: 
                        subprocess.run(['xdg-open', output_dir_abs], check=True)
                    self.log(f"Đã mở thư mục: {output_dir_abs}")
                except Exception as e:
                    self.log(f"Không thể tự động mở thư mục: {e}")
                    messagebox.showinfo("Thông báo", f"Không thể tự động mở thư mục. Vui lòng mở thủ công:\n{output_dir_abs}")
        else:
            self.log("Chuyển đổi thất bại!")

def create_sample_docx_with_image(): # Hàm này để tạo file mẫu, cần Pillow
    sample_docx_file = "mau_co_hinh.docx"
    # if os.path.exists(sample_docx_file): # Bỏ comment nếu không muốn tạo lại mỗi lần
    #     print(f"Tệp Word mẫu '{sample_docx_file}' đã tồn tại.")
    #     return sample_docx_file

    doc = Document()
    doc.add_heading('Tài liệu mẫu có hình ảnh', 0)
    
    p = doc.add_paragraph('Đây là một đoạn văn bản bình thường. ')
    p.add_run('Văn bản in đậm.').bold = True
    p.add_run(' Và đây là ')
    p.add_run('văn bản in nghiêng.').italic = True

    doc.add_paragraph('Một đoạn văn khác với & ký tự đặc biệt % $ # _ { } ~ ^ \\.')

    # Tạo và thêm ảnh PNG mẫu
    sample_image_path_png = "sample_image.png"
    try:
        from PIL import Image, ImageDraw
        img_png = Image.new('RGB', (200, 100), color = (255, 128, 128))
        d_png = ImageDraw.Draw(img_png)
        d_png.text((10,10), "Sample PNG", fill=(0,0,0))
        img_png.save(sample_image_path_png)
        print(f"Đã tạo tệp ảnh mẫu: {sample_image_path_png}")
        
        from docx.shared import Inches # Import ở đây khi cần
        doc.add_paragraph("Hình ảnh PNG inline bên dưới:")
        doc.add_picture(sample_image_path_png, width=Inches(2.5)) 
        doc.add_paragraph("Đoạn văn sau hình ảnh PNG.")
    except ImportError:
        print("PIL/Pillow không được cài đặt. Không thể tạo hoặc thêm ảnh mẫu PNG.")
    except Exception as e:
        print(f"Lỗi khi tạo hoặc thêm ảnh mẫu PNG: {e}")

    # Tạo và thêm ảnh SVG mẫu (Lưu ý: python-docx có thể không nhúng SVG trực tiếp vào DOCX)
    # Việc trích xuất SVG sẽ phụ thuộc vào cách Word lưu trữ nó.
    sample_image_path_svg = "sample_image.svg"
    try:
        svg_content = """<svg width="100" height="100" xmlns="http://www.w3.org/2000/svg">
                           <rect width="100" height="100" style="fill:rgb(0,0,255);stroke-width:3;stroke:rgb(0,0,0)" />
                           <text x="10" y="50" font-family="Verdana" font-size="15" fill="white">SVG</text>
                         </svg>"""
        with open(sample_image_path_svg, "w", encoding="utf-8") as f_svg:
            f_svg.write(svg_content)
        print(f"Đã tạo tệp ảnh mẫu: {sample_image_path_svg}")
        # doc.add_paragraph("Hình ảnh SVG (nếu Word hỗ trợ nhúng và python-docx trích xuất được):")
        # Hiện tại, không có cách chuẩn để thêm SVG trực tiếp vào DOCX bằng python-docx
        # mà đảm bảo nó sẽ được trích xuất như một image part.
    except Exception as e:
        print(f"Lỗi khi tạo ảnh mẫu SVG: {e}")


    doc.add_heading('Một tiêu đề cấp 2', level=2)
    doc.add_paragraph('Nội dung dưới tiêu đề cấp 2.')

    doc.save(sample_docx_file)
    print(f"Đã tạo/cập nhật tệp Word mẫu: {sample_docx_file}")
    return sample_docx_file


if __name__ == "__main__":
    # Bỏ comment dòng dưới nếu bạn muốn tạo file DOCX mẫu mỗi khi chạy script
    # create_sample_docx_with_image() 
    
    root = tk.Tk()
    app = WordToLaTeXConverterApp(root)
    root.mainloop()